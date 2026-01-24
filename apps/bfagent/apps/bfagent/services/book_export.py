"""
Book Export Service - DEPRECATED

================================================================================
DEPRECATED - DO NOT USE
================================================================================

This file has been deprecated and replaced by the consolidated Core Services.

Replacement: apps.core.services.export
Deprecated: 2025-12-07
Migration Phase: 5

This file is kept for reference only. All new code should use the replacement.

To migrate existing code, run:
    python manage.py migrate_to_core --apply

================================================================================
Export generated books to DOCX, PDF, EPUB formats
"""

import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

import markdown
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from ebooklib import epub
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

from apps.bfagent.models import BookProjects
from apps.core.services.storage import ContentStorageService

logger = logging.getLogger(__name__)


class BookExporter:
    """Service for exporting books to various formats"""

    def __init__(self):
        self.storage = ContentStorageService()

    def export_to_docx(self, project: BookProjects, output_path: Optional[Path] = None) -> Path:
        """
        Export book to Microsoft Word DOCX format

        Args:
            project: BookProjects instance
            output_path: Optional custom output path

        Returns:
            Path to created DOCX file
        """
        from django.utils.text import slugify

        project_slug = slugify(project.title)
        chapter_path = self.storage.get_chapter_path(project_slug)

        if output_path is None:
            export_path = self.storage.get_export_path(project_slug)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = export_path / f"{project_slug}_{timestamp}.docx"

        # Create document
        doc = Document()

        # Title page
        title = doc.add_heading(project.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        if project.genre:
            genre_para = doc.add_paragraph(f"Genre: {project.genre}")
            genre_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        doc.add_paragraph()

        author_para = doc.add_paragraph("Generated with BF Agent Book Generation System")
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        date_para = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # Get all chapters
        chapters = sorted(chapter_path.glob("chapter_*.md"))

        for chapter_file in chapters:
            content = chapter_file.read_text(encoding="utf-8")

            # Parse markdown and add to document
            lines = content.split("\n")

            for line in lines:
                line = line.strip()

                if not line:
                    doc.add_paragraph()
                    continue

                # Skip metadata section
                if line.startswith("## Metadata") or line.startswith("- **"):
                    continue
                if line == "---":
                    continue

                # Headings
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=2)
                elif line.startswith("#### "):
                    doc.add_heading(line[5:], level=3)
                else:
                    # Regular paragraph
                    if line:
                        doc.add_paragraph(line)

            doc.add_page_break()

        # Save
        doc.save(str(output_path))
        logger.info(f"Exported book to DOCX: {output_path}")

        return output_path

    def export_to_pdf(self, project: BookProjects, output_path: Optional[Path] = None) -> Path:
        """
        Export book to PDF format

        Args:
            project: BookProjects instance
            output_path: Optional custom output path

        Returns:
            Path to created PDF file
        """
        from django.utils.text import slugify

        project_slug = slugify(project.title)
        chapter_path = self.storage.get_chapter_path(project_slug)

        if output_path is None:
            export_path = self.storage.get_export_path(project_slug)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = export_path / f"{project_slug}_{timestamp}.pdf"

        # Create PDF
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )

        story = []
        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Heading1"],
            fontSize=24,
            textColor=RGBColor(0, 0, 0),
            spaceAfter=30,
            alignment=1,  # Center
        )

        chapter_style = ParagraphStyle(
            "ChapterTitle",
            parent=styles["Heading1"],
            fontSize=18,
            spaceAfter=12,
        )

        # Title page
        story.append(Spacer(1, 2 * inch))
        story.append(Paragraph(project.title, title_style))
        story.append(Spacer(1, 0.5 * inch))

        if project.genre:
            genre_style = ParagraphStyle("Genre", parent=styles["Normal"], alignment=1)
            story.append(Paragraph(f"Genre: {project.genre}", genre_style))

        story.append(PageBreak())

        # Chapters
        chapters = sorted(chapter_path.glob("chapter_*.md"))

        for chapter_file in chapters:
            content = chapter_file.read_text(encoding="utf-8")

            lines = content.split("\n")
            for line in lines:
                line = line.strip()

                if not line:
                    story.append(Spacer(1, 0.2 * inch))
                    continue

                # Skip metadata
                if line.startswith("## Metadata") or line.startswith("- **"):
                    continue
                if line == "---":
                    continue

                # Headings
                if line.startswith("### "):
                    story.append(Paragraph(line[4:], chapter_style))
                elif line.startswith("#### "):
                    story.append(Paragraph(line[5:], styles["Heading2"]))
                else:
                    if line:
                        story.append(Paragraph(line, styles["BodyText"]))
                        story.append(Spacer(1, 0.1 * inch))

            story.append(PageBreak())

        doc.build(story)
        logger.info(f"Exported book to PDF: {output_path}")

        return output_path

    def export_to_epub(self, project: BookProjects, output_path: Optional[Path] = None) -> Path:
        """
        Export book to EPUB format

        Args:
            project: BookProjects instance
            output_path: Optional custom output path

        Returns:
            Path to created EPUB file
        """
        from django.utils.text import slugify

        project_slug = slugify(project.title)
        chapter_path = self.storage.get_chapter_path(project_slug)

        if output_path is None:
            export_path = self.storage.get_export_path(project_slug)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = export_path / f"{project_slug}_{timestamp}.epub"

        # Create EPUB book
        book = epub.EpubBook()

        # Metadata
        book.set_identifier(f"bfagent-{project.id}")
        book.set_title(project.title)
        book.set_language("de")  # German
        book.add_author("BF Agent")

        if project.genre:
            book.add_metadata("DC", "subject", project.genre)

        # Chapters
        chapters_list = []
        chapters = sorted(chapter_path.glob("chapter_*.md"))

        for idx, chapter_file in enumerate(chapters, 1):
            content = chapter_file.read_text(encoding="utf-8")

            # Convert markdown to HTML
            md = markdown.Markdown(extensions=["extra"])
            html_content = md.convert(content)

            # Create EPUB chapter
            c = epub.EpubHtml(title=f"Chapter {idx}", file_name=f"chap_{idx:02d}.xhtml", lang="de")

            c.content = f"<h1>Chapter {idx}</h1>{html_content}"

            book.add_item(c)
            chapters_list.append(c)

        # Table of contents
        book.toc = tuple(chapters_list)

        # Navigation
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        # Spine
        book.spine = ["nav"] + chapters_list

        # Write EPUB file
        epub.write_epub(str(output_path), book)
        logger.info(f"Exported book to EPUB: {output_path}")

        return output_path

    def export_all_formats(self, project: BookProjects) -> dict:
        """
        Export book to all supported formats

        Returns:
            Dict with paths to all exported files
        """
        results = {}

        try:
            results["docx"] = str(self.export_to_docx(project))
        except Exception as e:
            logger.error(f"DOCX export failed: {e}")
            results["docx"] = None

        try:
            results["pdf"] = str(self.export_to_pdf(project))
        except Exception as e:
            logger.error(f"PDF export failed: {e}")
            results["pdf"] = None

        try:
            results["epub"] = str(self.export_to_epub(project))
        except Exception as e:
            logger.error(f"EPUB export failed: {e}")
            results["epub"] = None

        return results
