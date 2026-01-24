"""
Export Service - Concrete Exporters

Implementations for various export formats:
- DOCXExporter: Microsoft Word documents
- PDFExporter: PDF documents
- EPUBExporter: E-book format
- MarkdownExporter: Markdown files
- HTMLExporter: HTML documents
- JSONExporter: JSON data
- CSVExporter: CSV data

Part of the consolidated Core Export Service.
"""

import logging
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

try:
    from .base import BaseExporter, ContentConverter
    from .exceptions import ExportException, MissingDependencyError, WriteError, wrap_library_error
    from .models import (
        BookContent,
        ChapterContent,
        DocumentMetadata,
        DOCXExportConfig,
        EPUBExportConfig,
        ExportConfig,
        ExportFormat,
        ExportResult,
        MarkdownExportConfig,
        PDFExportConfig,
    )
except ImportError:
    from base import BaseExporter, ContentConverter
    from exceptions import ExportException, MissingDependencyError, WriteError, wrap_library_error
    from models import (
        BookContent,
        ChapterContent,
        DocumentMetadata,
        DOCXExportConfig,
        EPUBExportConfig,
        ExportConfig,
        ExportFormat,
        ExportResult,
        MarkdownExportConfig,
        PDFExportConfig,
    )


logger = logging.getLogger(__name__)


# =============================================================================
# DOCX Exporter
# =============================================================================


class DOCXExporter(BaseExporter):
    """
    Microsoft Word DOCX exporter.

    Uses python-docx library to create Word documents.

    Example:
        exporter = DOCXExporter()
        result = exporter.export(
            content="# Chapter 1\\n\\nContent here...",
            metadata=DocumentMetadata(title="My Book")
        )
    """

    format = ExportFormat.DOCX
    file_extension = "docx"

    def _export(
        self, content: Any, output_path: Path, metadata: Optional[DocumentMetadata] = None
    ) -> ExportResult:
        """Export content to DOCX."""
        result = ExportResult(format=self.format)

        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches, Pt
        except ImportError:
            raise MissingDependencyError(
                "python-docx", format="docx", install_cmd="pip install python-docx"
            )

        doc = Document()

        # Add title if metadata provided
        if metadata and metadata.title:
            title = doc.add_heading(metadata.title, level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if metadata.author:
                author_para = doc.add_paragraph(f"By {metadata.author}")
                author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_page_break()

        # Handle different content types
        if isinstance(content, BookContent):
            self._add_book_content(doc, content)
        elif isinstance(content, str):
            self._add_markdown_content(doc, content)
        elif isinstance(content, list):
            for item in content:
                if isinstance(item, ChapterContent):
                    self._add_chapter(doc, item)
                elif isinstance(item, str):
                    self._add_markdown_content(doc, item)

        # Save document
        doc.save(str(output_path))
        result.metadata = metadata

        return result

    def _add_book_content(self, doc, book: BookContent) -> None:
        """Add complete book content."""
        for chapter in book.chapters:
            self._add_chapter(doc, chapter)

    def _add_chapter(self, doc, chapter: ChapterContent) -> None:
        """Add single chapter."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Chapter heading
        title = chapter.title or f"Chapter {chapter.number}"
        doc.add_heading(title, level=1)

        # Content
        self._add_markdown_content(doc, chapter.content)

        # Page break after chapter
        doc.add_page_break()

    def _add_markdown_content(self, doc, markdown_text: str) -> None:
        """Convert markdown to docx paragraphs."""
        lines = markdown_text.split("\n")

        for line in lines:
            line = line.strip()

            if not line:
                doc.add_paragraph()
                continue

            # Skip frontmatter markers
            if line == "---":
                continue

            # Headings
            if line.startswith("#### "):
                doc.add_heading(line[5:], level=4)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            else:
                # Regular paragraph
                doc.add_paragraph(line)


# =============================================================================
# PDF Exporter
# =============================================================================


class PDFExporter(BaseExporter):
    """
    PDF document exporter.

    Uses ReportLab library to create PDF documents.

    Example:
        exporter = PDFExporter()
        result = exporter.export(
            content="# Chapter 1\\n\\nContent here...",
            metadata=DocumentMetadata(title="My Book")
        )
    """

    format = ExportFormat.PDF
    file_extension = "pdf"

    def _export(
        self, content: Any, output_path: Path, metadata: Optional[DocumentMetadata] = None
    ) -> ExportResult:
        """Export content to PDF."""
        result = ExportResult(format=self.format)

        try:
            from reportlab.lib.pagesizes import A4, letter
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import inch
            from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer
        except ImportError:
            raise MissingDependencyError(
                "reportlab", format="pdf", install_cmd="pip install reportlab"
            )

        # Get PDF options
        pdf_config = self.config.pdf_options or {}
        page_size = letter if pdf_config.get("page_size", "letter") == "letter" else A4
        margin = pdf_config.get("margin_inches", 1.0)

        # Create document
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=page_size,
            leftMargin=margin * inch,
            rightMargin=margin * inch,
            topMargin=margin * inch,
            bottomMargin=margin * inch,
        )

        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "CustomTitle", parent=styles["Title"], fontSize=24, spaceAfter=30
        )
        heading_style = ParagraphStyle(
            "CustomHeading", parent=styles["Heading1"], fontSize=18, spaceAfter=12, spaceBefore=24
        )

        story = []

        # Title page
        if metadata and metadata.title:
            story.append(Spacer(1, 2 * inch))
            story.append(Paragraph(metadata.title, title_style))
            if metadata.author:
                story.append(Spacer(1, 0.5 * inch))
                story.append(Paragraph(f"By {metadata.author}", styles["Normal"]))
            story.append(PageBreak())

        # Content
        if isinstance(content, BookContent):
            for chapter in content.chapters:
                story.extend(self._chapter_to_story(chapter, styles, heading_style))
        elif isinstance(content, str):
            story.extend(self._markdown_to_story(content, styles, heading_style))

        # Build PDF
        doc.build(story)
        result.metadata = metadata

        return result

    def _chapter_to_story(self, chapter: ChapterContent, styles, heading_style) -> List:
        """Convert chapter to story elements."""
        from reportlab.lib.units import inch
        from reportlab.platypus import PageBreak, Paragraph, Spacer

        story = []

        # Chapter heading
        title = chapter.title or f"Chapter {chapter.number}"
        story.append(Paragraph(title, heading_style))
        story.append(Spacer(1, 0.2 * inch))

        # Content
        story.extend(self._markdown_to_story(chapter.content, styles, heading_style))
        story.append(PageBreak())

        return story

    def _markdown_to_story(self, markdown_text: str, styles, heading_style) -> List:
        """Convert markdown to story elements."""
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, Spacer

        story = []
        lines = markdown_text.split("\n")

        for line in lines:
            line = line.strip()

            if not line:
                story.append(Spacer(1, 0.1 * inch))
                continue

            # Skip frontmatter
            if line == "---" or line.startswith("- **"):
                continue

            # Headings
            if line.startswith("# "):
                story.append(Paragraph(line[2:], heading_style))
            elif line.startswith("## "):
                story.append(Paragraph(line[3:], styles["Heading2"]))
            elif line.startswith("### "):
                story.append(Paragraph(line[4:], styles["Heading3"]))
            else:
                story.append(Paragraph(line, styles["BodyText"]))

        return story


# =============================================================================
# EPUB Exporter
# =============================================================================


class EPUBExporter(BaseExporter):
    """
    EPUB e-book exporter.

    Uses ebooklib library to create EPUB files.

    Example:
        exporter = EPUBExporter()
        result = exporter.export(
            content=book_content,
            metadata=DocumentMetadata(title="My Book")
        )
    """

    format = ExportFormat.EPUB
    file_extension = "epub"

    def _export(
        self, content: Any, output_path: Path, metadata: Optional[DocumentMetadata] = None
    ) -> ExportResult:
        """Export content to EPUB."""
        result = ExportResult(format=self.format)

        try:
            from ebooklib import epub
        except ImportError:
            raise MissingDependencyError(
                "ebooklib", format="epub", install_cmd="pip install ebooklib"
            )

        book = epub.EpubBook()

        # Metadata
        title = metadata.title if metadata else "Untitled"
        book.set_identifier(f"export-{datetime.now().timestamp()}")
        book.set_title(title)
        book.set_language(metadata.language if metadata else "en")

        if metadata and metadata.author:
            book.add_author(metadata.author)

        # Create chapters
        chapters_list = []

        if isinstance(content, BookContent):
            for idx, chapter in enumerate(content.chapters, 1):
                epub_chapter = self._create_epub_chapter(chapter, idx)
                book.add_item(epub_chapter)
                chapters_list.append(epub_chapter)
        elif isinstance(content, str):
            epub_chapter = self._create_epub_chapter_from_text(content, 1)
            book.add_item(epub_chapter)
            chapters_list.append(epub_chapter)

        # Table of contents
        book.toc = tuple(chapters_list)

        # Navigation
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        # Spine
        book.spine = ["nav"] + chapters_list

        # Write EPUB
        epub.write_epub(str(output_path), book)
        result.metadata = metadata

        return result

    def _create_epub_chapter(self, chapter: ChapterContent, idx: int):
        """Create EPUB chapter from ChapterContent."""
        from ebooklib import epub

        title = chapter.title or f"Chapter {chapter.number}"

        c = epub.EpubHtml(title=title, file_name=f"chap_{idx:02d}.xhtml", lang="en")

        html_content = ContentConverter.markdown_to_html(chapter.content)
        c.content = f"<h1>{title}</h1>{html_content}"

        return c

    def _create_epub_chapter_from_text(self, text: str, idx: int):
        """Create EPUB chapter from text."""
        from ebooklib import epub

        c = epub.EpubHtml(title=f"Chapter {idx}", file_name=f"chap_{idx:02d}.xhtml", lang="en")

        html_content = ContentConverter.markdown_to_html(text)
        c.content = html_content

        return c


# =============================================================================
# Markdown Exporter
# =============================================================================


class MarkdownExporter(BaseExporter):
    """
    Markdown file exporter.

    Exports content to Markdown files with optional YAML frontmatter.

    Example:
        exporter = MarkdownExporter()
        result = exporter.export(
            content={"title": "My Doc", "content": "Hello world"},
            metadata=DocumentMetadata(title="My Doc")
        )
    """

    format = ExportFormat.MARKDOWN
    file_extension = "md"

    def _export(
        self, content: Any, output_path: Path, metadata: Optional[DocumentMetadata] = None
    ) -> ExportResult:
        """Export content to Markdown."""
        result = ExportResult(format=self.format)

        # Convert content to markdown
        if isinstance(content, str):
            md_content = content
        elif isinstance(content, dict):
            md_content = content.get("content", "")
            if "title" in content and not md_content.startswith("#"):
                md_content = f"# {content['title']}\n\n{md_content}"
        elif isinstance(content, BookContent):
            md_content = self._book_to_markdown(content)
        else:
            md_content = str(content)

        # Add frontmatter
        if self.config.add_metadata and metadata:
            md_content = ContentConverter.add_yaml_frontmatter(md_content, metadata.to_dict())

        # Write file
        output_path.write_text(md_content, encoding="utf-8")
        result.metadata = metadata

        return result

    def _book_to_markdown(self, book: BookContent) -> str:
        """Convert book to markdown."""
        lines = [f"# {book.title}", ""]

        if book.author:
            lines.extend([f"*By {book.author}*", ""])

        for chapter in book.chapters:
            title = chapter.title or f"Chapter {chapter.number}"
            lines.extend([f"## {title}", "", chapter.content, ""])

        return "\n".join(lines)


# =============================================================================
# HTML Exporter
# =============================================================================


class HTMLExporter(BaseExporter):
    """
    HTML document exporter.

    Exports content to HTML files with optional styling.

    Example:
        exporter = HTMLExporter()
        result = exporter.export(
            content="# My Document\\n\\nContent here.",
            metadata=DocumentMetadata(title="My Doc")
        )
    """

    format = ExportFormat.HTML
    file_extension = "html"

    def _export(
        self, content: Any, output_path: Path, metadata: Optional[DocumentMetadata] = None
    ) -> ExportResult:
        """Export content to HTML."""
        result = ExportResult(format=self.format)

        # Convert content
        if isinstance(content, str):
            body_content = ContentConverter.markdown_to_html(content)
        elif isinstance(content, dict):
            body_content = ContentConverter.markdown_to_html(content.get("content", ""))
        else:
            body_content = f"<p>{content}</p>"

        # Build HTML document
        title = metadata.title if metadata else "Document"

        html_options = self.config.html_options or {}
        css = html_options.get("css", self._default_css())

        html_content = f"""<!DOCTYPE html>
<html lang="{metadata.language if metadata else 'en'}">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>{css}</style>
</head>
<body>
    <article>
        {body_content}
    </article>
</body>
</html>"""

        output_path.write_text(html_content, encoding="utf-8")
        result.metadata = metadata

        return result

    def _default_css(self) -> str:
        """Get default CSS styling."""
        return """
        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 2rem;
            color: #333;
        }
        h1, h2, h3 { margin-top: 1.5em; }
        h1 { border-bottom: 2px solid #eee; padding-bottom: 0.5em; }
        code { background: #f4f4f4; padding: 0.2em 0.4em; border-radius: 3px; }
        pre { background: #f4f4f4; padding: 1em; overflow-x: auto; }
        blockquote { border-left: 4px solid #ddd; margin-left: 0; padding-left: 1em; color: #666; }
        """


# =============================================================================
# JSON Exporter
# =============================================================================


class JSONExporter(BaseExporter):
    """
    JSON data exporter.

    Exports data to JSON files.

    Example:
        exporter = JSONExporter()
        result = exporter.export(
            content={"key": "value", "list": [1, 2, 3]}
        )
    """

    format = ExportFormat.JSON
    file_extension = "json"

    def _export(
        self, content: Any, output_path: Path, metadata: Optional[DocumentMetadata] = None
    ) -> ExportResult:
        """Export content to JSON."""
        result = ExportResult(format=self.format)

        json_content = ContentConverter.dict_to_json(content)
        output_path.write_text(json_content, encoding="utf-8")
        result.metadata = metadata

        return result


# =============================================================================
# CSV Exporter
# =============================================================================


class CSVExporter(BaseExporter):
    """
    CSV data exporter.

    Exports tabular data to CSV files.

    Example:
        exporter = CSVExporter()
        result = exporter.export(
            content=[
                {"name": "Alice", "age": 30},
                {"name": "Bob", "age": 25}
            ]
        )
    """

    format = ExportFormat.CSV
    file_extension = "csv"

    def validate_content(self, content: Any) -> bool:
        """Validate content is list of dicts."""
        if not isinstance(content, list):
            return False
        if not content:
            return False
        return all(isinstance(item, dict) for item in content)

    def _export(
        self, content: Any, output_path: Path, metadata: Optional[DocumentMetadata] = None
    ) -> ExportResult:
        """Export content to CSV."""
        result = ExportResult(format=self.format)

        csv_content = ContentConverter.dict_to_csv(content)
        output_path.write_text(csv_content, encoding="utf-8")
        result.metadata = metadata

        return result


# =============================================================================
# Factory Function
# =============================================================================


def create_exporter(
    format: Union[str, ExportFormat], config: Optional[ExportConfig] = None
) -> BaseExporter:
    """
    Factory function to create exporters.

    Args:
        format: Export format
        config: Export configuration

    Returns:
        Exporter instance

    Example:
        exporter = create_exporter("docx")
        result = exporter.export(content, metadata=metadata)
    """
    try:
        from .exceptions import UnsupportedFormatError
    except ImportError:
        from exceptions import UnsupportedFormatError

    if isinstance(format, str):
        try:
            format = ExportFormat(format.lower())
        except ValueError:
            pass

    exporters = {
        ExportFormat.DOCX: DOCXExporter,
        ExportFormat.PDF: PDFExporter,
        ExportFormat.EPUB: EPUBExporter,
        ExportFormat.MARKDOWN: MarkdownExporter,
        ExportFormat.HTML: HTMLExporter,
        ExportFormat.JSON: JSONExporter,
        ExportFormat.CSV: CSVExporter,
        "md": MarkdownExporter,
    }

    exporter_class = exporters.get(format)
    if exporter_class is None:
        raise UnsupportedFormatError(str(format), supported=[f.value for f in ExportFormat])

    return exporter_class(config)
