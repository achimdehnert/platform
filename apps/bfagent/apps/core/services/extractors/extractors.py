"""
File Extractors - Concrete Implementations

Implementations for various file types:
- PDFExtractor: PDF documents
- DOCXExtractor: Word documents
- PPTXExtractor: PowerPoint presentations
- ExcelExtractor: Excel spreadsheets
- CSVExtractor: CSV files
- JSONExtractor: JSON files
- TextExtractor: Plain text/Markdown
- ImageExtractor: OCR for images

Part of the consolidated Core File Extractor Service.
"""

import csv
import io
import json
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

try:
    from .base import BaseExtractor, TableParser, TextCleaner
    from .exceptions import (
        ExtractionError,
        FileReadError,
        MissingDependencyError,
        OCRError,
        TableExtractionError,
        TextExtractionError,
        wrap_library_error,
    )
    from .models import (
        ExtractedSlide,
        ExtractedTable,
        ExtractedText,
        ExtractionResult,
        ExtractorConfig,
        FileMetadata,
        FileType,
        detect_file_type,
    )
except ImportError:
    from base import BaseExtractor, TableParser, TextCleaner
    from exceptions import (
        ExtractionError,
        FileReadError,
        MissingDependencyError,
        OCRError,
        TableExtractionError,
        TextExtractionError,
        wrap_library_error,
    )
    from models import (
        ExtractedSlide,
        ExtractedTable,
        ExtractedText,
        ExtractionResult,
        ExtractorConfig,
        FileMetadata,
        FileType,
        detect_file_type,
    )


logger = logging.getLogger(__name__)


# =============================================================================
# PDF Extractor
# =============================================================================


class PDFExtractor(BaseExtractor):
    """
    PDF document extractor.

    Uses pdfplumber for text extraction, with optional OCR fallback.

    Features:
        - Text extraction with layout preservation
        - Table extraction
        - Metadata extraction
        - Page range selection
        - OCR for scanned documents (optional)

    Example:
        extractor = PDFExtractor()
        result = extractor.extract("document.pdf")
        print(result.text)
    """

    supported_types = [FileType.PDF]
    file_extensions = [".pdf"]

    def _extract(self, file_path: Path, config: ExtractorConfig) -> ExtractionResult:
        """Extract content from PDF."""
        result = ExtractionResult(file_type=FileType.PDF)

        try:
            import pdfplumber
        except ImportError:
            raise MissingDependencyError(
                "pdfplumber", file_type="pdf", install_cmd="pip install pdfplumber"
            )

        try:
            with pdfplumber.open(str(file_path)) as pdf:
                # Metadata
                result.metadata = FileMetadata(
                    filename=file_path.name,
                    file_type=FileType.PDF,
                    page_count=len(pdf.pages),
                    title=pdf.metadata.get("Title", ""),
                    author=pdf.metadata.get("Author", ""),
                )

                # Determine pages to extract
                pages_to_extract = config.page_range or range(len(pdf.pages))

                for page_num in pages_to_extract:
                    if page_num >= len(pdf.pages):
                        continue

                    page = pdf.pages[page_num]

                    # Extract text
                    text = page.extract_text() or ""

                    if text.strip():
                        result.texts.append(
                            ExtractedText(text=TextCleaner.clean(text), page_number=page_num + 1)
                        )
                    elif config.ocr_enabled:
                        # Try OCR for this page
                        ocr_text = self._ocr_page(page, config)
                        if ocr_text:
                            result.texts.append(
                                ExtractedText(
                                    text=ocr_text,
                                    page_number=page_num + 1,
                                    formatting={"source": "ocr"},
                                )
                            )

                    # Extract tables
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            headers = [str(h or "") for h in table[0]]
                            rows = TableParser.rows_to_dicts(table[1:], headers)
                            result.tables.append(
                                ExtractedTable(headers=headers, rows=rows, page_number=page_num + 1)
                            )

                # Combine all text
                result.raw_content = "\n\n".join(t.text for t in result.texts)

        except Exception as e:
            raise wrap_library_error(e, "pdf", "pdfplumber")

        return result

    def _ocr_page(self, page, config: ExtractorConfig) -> str:
        """Perform OCR on a page."""
        try:
            import pytesseract
            from PIL import Image
        except ImportError:
            logger.warning("OCR not available: install pytesseract and pillow")
            return ""

        try:
            # Convert page to image
            image = page.to_image(resolution=300)
            pil_image = image.original

            # OCR
            text = pytesseract.image_to_string(pil_image, lang=config.ocr_language)

            return TextCleaner.clean(text)
        except Exception as e:
            logger.warning(f"OCR failed: {e}")
            return ""


# =============================================================================
# DOCX Extractor
# =============================================================================


class DOCXExtractor(BaseExtractor):
    """
    Microsoft Word DOCX extractor.

    Uses python-docx for text extraction.

    Features:
        - Paragraph text extraction
        - Table extraction
        - Metadata extraction
        - Style preservation (optional)

    Example:
        extractor = DOCXExtractor()
        result = extractor.extract("document.docx")
    """

    supported_types = [FileType.DOCX]
    file_extensions = [".docx"]

    def _extract(self, file_path: Path, config: ExtractorConfig) -> ExtractionResult:
        """Extract content from DOCX."""
        result = ExtractionResult(file_type=FileType.DOCX)

        try:
            from docx import Document
        except ImportError:
            raise MissingDependencyError(
                "python-docx", file_type="docx", install_cmd="pip install python-docx"
            )

        try:
            doc = Document(str(file_path))

            # Metadata
            core_props = doc.core_properties
            result.metadata = FileMetadata(
                filename=file_path.name,
                file_type=FileType.DOCX,
                title=core_props.title or "",
                author=core_props.author or "",
                modified_at=core_props.modified,
            )

            # Extract paragraphs
            texts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    texts.append(para.text)

            result.texts.append(ExtractedText(text="\n\n".join(texts)))

            # Extract tables
            for table in doc.tables:
                headers = []
                rows = []

                for i, row in enumerate(table.rows):
                    row_data = [cell.text for cell in row.cells]

                    if i == 0:
                        headers = row_data
                    else:
                        row_dict = {
                            headers[j]: row_data[j] if j < len(row_data) else ""
                            for j in range(len(headers))
                        }
                        rows.append(row_dict)

                if headers:
                    result.tables.append(ExtractedTable(headers=headers, rows=rows))

            # Calculate word count
            result.metadata.word_count = sum(len(t.text.split()) for t in result.texts)

        except Exception as e:
            raise wrap_library_error(e, "docx", "python-docx")

        return result


# =============================================================================
# PPTX Extractor
# =============================================================================


class PPTXExtractor(BaseExtractor):
    """
    PowerPoint PPTX extractor.

    Extracts text from slides using XML parsing.
    Compatible with medtrans PPTXExtractor.

    Features:
        - Slide text extraction
        - Speaker notes extraction
        - Preserves slide structure
        - Shape-level text IDs

    Example:
        extractor = PPTXExtractor()
        result = extractor.extract("presentation.pptx")
        for slide in result.slides:
            print(f"Slide {slide.slide_number}: {slide.title}")
    """

    supported_types = [FileType.PPTX]
    file_extensions = [".pptx"]

    NAMESPACE_MAP = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "p": "http://schemas.openxmlformats.org/presentationml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }

    def _extract(self, file_path: Path, config: ExtractorConfig) -> ExtractionResult:
        """Extract content from PPTX."""
        import re
        import tempfile
        import xml.etree.ElementTree as ET
        import zipfile

        result = ExtractionResult(file_type=FileType.PPTX)

        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_path = Path(temp_dir)

                # Extract PPTX (it's a ZIP file)
                with zipfile.ZipFile(str(file_path), "r") as zip_ref:
                    zip_ref.extractall(temp_path)

                # Find slide XML files
                slides_dir = temp_path / "ppt" / "slides"
                if not slides_dir.exists():
                    result.warnings.append("No slides directory found")
                    return result

                slide_files = sorted(slides_dir.glob("slide*.xml"))
                result.metadata = FileMetadata(
                    filename=file_path.name, file_type=FileType.PPTX, page_count=len(slide_files)
                )

                # Process each slide
                for slide_num, slide_file in enumerate(slide_files, 1):
                    slide_data = self._extract_slide(slide_file, slide_num)
                    result.slides.append(slide_data)

                    # Also add to texts for unified access
                    if slide_data.content:
                        result.texts.append(
                            ExtractedText(text=slide_data.content, page_number=slide_num)
                        )

        except Exception as e:
            raise wrap_library_error(e, "pptx", "zipfile")

        return result

    def _extract_slide(self, slide_file: Path, slide_number: int) -> ExtractedSlide:
        """Extract content from a single slide."""
        import re
        import xml.etree.ElementTree as ET

        slide = ExtractedSlide(slide_number=slide_number)

        try:
            # Register namespaces
            for prefix, uri in self.NAMESPACE_MAP.items():
                ET.register_namespace(prefix, uri)

            tree = ET.parse(str(slide_file))
            root = tree.getroot()

            # Find all text elements
            text_elements = root.findall(".//a:t", self.NAMESPACE_MAP)

            texts = []
            for idx, elem in enumerate(text_elements):
                if elem.text and elem.text.strip():
                    text = elem.text.strip()

                    # Skip very short texts
                    if len(text) < 2:
                        continue

                    # Skip non-text content
                    if not re.search(r"[a-zA-ZäöüÄÖÜß]", text):
                        continue

                    texts.append(text)
                    slide.texts.append({"id": f"slide_{slide_number}_text_{idx + 1}", "text": text})

            # First text is usually the title
            if texts:
                slide.title = texts[0]
                slide.content = "\n".join(texts)

        except Exception as e:
            logger.warning(f"Error extracting slide {slide_number}: {e}")

        return slide


# =============================================================================
# Excel Extractor
# =============================================================================


class ExcelExtractor(BaseExtractor):
    """
    Excel spreadsheet extractor.

    Uses openpyxl for XLSX files.

    Features:
        - Multiple sheet support
        - Header detection
        - Data type preservation
        - Sheet selection

    Example:
        extractor = ExcelExtractor()
        result = extractor.extract("data.xlsx")
        for table in result.tables:
            print(f"Sheet: {table.sheet_name}, Rows: {table.row_count}")
    """

    supported_types = [FileType.XLSX, FileType.XLS]
    file_extensions = [".xlsx", ".xls"]

    def _extract(self, file_path: Path, config: ExtractorConfig) -> ExtractionResult:
        """Extract content from Excel."""
        result = ExtractionResult(file_type=FileType.XLSX)

        try:
            import openpyxl
        except ImportError:
            raise MissingDependencyError(
                "openpyxl", file_type="xlsx", install_cmd="pip install openpyxl"
            )

        try:
            wb = openpyxl.load_workbook(str(file_path), data_only=True)

            result.metadata = FileMetadata(
                filename=file_path.name,
                file_type=FileType.XLSX,
                custom={"sheet_count": len(wb.sheetnames)},
            )

            # Determine which sheets to process
            sheets = config.sheet_names or wb.sheetnames

            for sheet_name in sheets:
                if sheet_name not in wb.sheetnames:
                    result.warnings.append(f"Sheet not found: {sheet_name}")
                    continue

                ws = wb[sheet_name]

                # Extract data
                rows = list(ws.iter_rows(values_only=True))
                if not rows:
                    continue

                # First row as headers
                headers = [str(h or f"col_{i}") for i, h in enumerate(rows[0])]

                # Convert to dicts
                data_rows = []
                for row in rows[1:]:
                    row_dict = {
                        headers[i]: row[i] if i < len(row) else None for i in range(len(headers))
                    }
                    data_rows.append(row_dict)

                result.tables.append(
                    ExtractedTable(headers=headers, rows=data_rows, sheet_name=sheet_name)
                )

        except Exception as e:
            raise wrap_library_error(e, "xlsx", "openpyxl")

        return result


# =============================================================================
# CSV Extractor
# =============================================================================


class CSVExtractor(BaseExtractor):
    """
    CSV file extractor.

    Features:
        - Delimiter detection
        - Header handling
        - Encoding support

    Example:
        extractor = CSVExtractor()
        result = extractor.extract("data.csv")
    """

    supported_types = [FileType.CSV]
    file_extensions = [".csv"]

    def _extract(self, file_path: Path, config: ExtractorConfig) -> ExtractionResult:
        """Extract content from CSV."""
        result = ExtractionResult(file_type=FileType.CSV)

        try:
            with open(file_path, "r", encoding=config.encoding, newline="") as f:
                # Detect delimiter
                sample = f.read(4096)
                f.seek(0)

                try:
                    dialect = csv.Sniffer().sniff(sample)
                except csv.Error:
                    dialect = csv.excel

                reader = csv.DictReader(f, dialect=dialect)

                headers = reader.fieldnames or []
                rows = list(reader)

                result.tables.append(ExtractedTable(headers=headers, rows=rows))

                result.metadata = FileMetadata(
                    filename=file_path.name,
                    file_type=FileType.CSV,
                    custom={"row_count": len(rows), "column_count": len(headers)},
                )

        except UnicodeDecodeError as e:
            raise FileReadError(str(file_path), f"encoding error: {e}")
        except Exception as e:
            raise wrap_library_error(e, "csv", "csv")

        return result


# =============================================================================
# JSON Extractor
# =============================================================================


class JSONExtractor(BaseExtractor):
    """
    JSON file extractor.

    Features:
        - Nested structure handling
        - Array to table conversion
        - Pretty text output

    Example:
        extractor = JSONExtractor()
        result = extractor.extract("data.json")
    """

    supported_types = [FileType.JSON]
    file_extensions = [".json"]

    def _extract(self, file_path: Path, config: ExtractorConfig) -> ExtractionResult:
        """Extract content from JSON."""
        result = ExtractionResult(file_type=FileType.JSON)

        try:
            with open(file_path, "r", encoding=config.encoding) as f:
                data = json.load(f)

            # Store raw content
            result.raw_content = json.dumps(data, indent=2, ensure_ascii=False)

            # Convert arrays to tables
            if isinstance(data, list) and data and isinstance(data[0], dict):
                headers = list(data[0].keys())
                result.tables.append(ExtractedTable(headers=headers, rows=data))

            result.metadata = FileMetadata(filename=file_path.name, file_type=FileType.JSON)

        except json.JSONDecodeError as e:
            raise FileReadError(str(file_path), f"invalid JSON: {e}")
        except Exception as e:
            raise wrap_library_error(e, "json", "json")

        return result


# =============================================================================
# Text/Markdown Extractor
# =============================================================================


class TextExtractor(BaseExtractor):
    """
    Plain text and Markdown extractor.

    Features:
        - Encoding detection
        - Markdown metadata extraction
        - Line-by-line processing

    Example:
        extractor = TextExtractor()
        result = extractor.extract("readme.md")
    """

    supported_types = [FileType.TXT, FileType.MARKDOWN, FileType.HTML]
    file_extensions = [".txt", ".md", ".markdown", ".html", ".htm"]

    def _extract(self, file_path: Path, config: ExtractorConfig) -> ExtractionResult:
        """Extract content from text file."""
        file_type = detect_file_type(file_path) or FileType.TXT
        result = ExtractionResult(file_type=file_type)

        try:
            with open(file_path, "r", encoding=config.encoding) as f:
                content = f.read()

            result.raw_content = content
            result.texts.append(ExtractedText(text=content))

            # Extract markdown metadata if present
            if file_type == FileType.MARKDOWN:
                metadata = self._extract_frontmatter(content)
                result.metadata = FileMetadata(
                    filename=file_path.name,
                    file_type=file_type,
                    title=metadata.get("title", ""),
                    author=metadata.get("author", ""),
                    custom=metadata,
                )
            else:
                result.metadata = FileMetadata(
                    filename=file_path.name, file_type=file_type, word_count=len(content.split())
                )

        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, "r", encoding="latin-1") as f:
                    content = f.read()
                result.raw_content = content
                result.texts.append(ExtractedText(text=content))
                result.warnings.append("File read with latin-1 encoding")
            except Exception as e:
                raise FileReadError(str(file_path), f"encoding error: {e}")

        return result

    def _extract_frontmatter(self, content: str) -> Dict[str, Any]:
        """Extract YAML frontmatter from markdown."""
        import re

        match = re.match(r"^---\s*\n(.+?)\n---\s*\n", content, re.DOTALL)
        if not match:
            return {}

        try:
            import yaml

            return yaml.safe_load(match.group(1)) or {}
        except:
            return {}


# =============================================================================
# Image Extractor (OCR)
# =============================================================================


class ImageExtractor(BaseExtractor):
    """
    Image OCR extractor.

    Uses pytesseract for text extraction from images.

    Features:
        - Multiple image format support
        - Language selection
        - Preprocessing options

    Example:
        extractor = ImageExtractor()
        result = extractor.extract("scan.png")
    """

    supported_types = [FileType.PNG, FileType.JPG, FileType.JPEG, FileType.TIFF]
    file_extensions = [".png", ".jpg", ".jpeg", ".tiff", ".tif"]

    def _extract(self, file_path: Path, config: ExtractorConfig) -> ExtractionResult:
        """Extract text from image using OCR."""
        result = ExtractionResult(file_type=detect_file_type(file_path))

        try:
            import pytesseract
            from PIL import Image
        except ImportError:
            raise MissingDependencyError(
                "pytesseract and pillow",
                file_type="image",
                install_cmd="pip install pytesseract pillow",
            )

        try:
            image = Image.open(str(file_path))

            text = pytesseract.image_to_string(image, lang=config.ocr_language)

            result.raw_content = TextCleaner.clean(text)
            result.texts.append(
                ExtractedText(text=result.raw_content, formatting={"source": "ocr"})
            )

            result.metadata = FileMetadata(
                filename=file_path.name,
                file_type=result.file_type,
                custom={"width": image.width, "height": image.height, "format": image.format},
            )

        except Exception as e:
            raise wrap_library_error(e, "image", "pytesseract")

        return result


# =============================================================================
# Factory Function
# =============================================================================


def create_extractor(
    file_type: Union[str, FileType, Path], config: Optional[ExtractorConfig] = None
) -> BaseExtractor:
    """
    Factory function to create appropriate extractor.

    Args:
        file_type: File type, extension, or path
        config: Extractor configuration

    Returns:
        Appropriate extractor instance

    Example:
        extractor = create_extractor("pdf")
        # or
        extractor = create_extractor(Path("document.pdf"))
    """
    try:
        from .exceptions import UnsupportedFormatError
    except ImportError:
        from exceptions import UnsupportedFormatError

    # Handle Path
    if isinstance(file_type, Path):
        file_type = detect_file_type(file_type)

    # Handle string
    if isinstance(file_type, str):
        file_type = FileType.from_extension(file_type) or file_type

    extractors = {
        FileType.PDF: PDFExtractor,
        FileType.DOCX: DOCXExtractor,
        FileType.PPTX: PPTXExtractor,
        FileType.XLSX: ExcelExtractor,
        FileType.XLS: ExcelExtractor,
        FileType.CSV: CSVExtractor,
        FileType.JSON: JSONExtractor,
        FileType.TXT: TextExtractor,
        FileType.MARKDOWN: TextExtractor,
        FileType.HTML: TextExtractor,
        FileType.PNG: ImageExtractor,
        FileType.JPG: ImageExtractor,
        FileType.JPEG: ImageExtractor,
        FileType.TIFF: ImageExtractor,
    }

    extractor_class = extractors.get(file_type)
    if extractor_class is None:
        supported = [t.value for t in extractors.keys()]
        raise UnsupportedFormatError(str(file_type), supported=supported)

    return extractor_class(config)


def extract_file(
    file_path: Union[str, Path], config: Optional[ExtractorConfig] = None
) -> ExtractionResult:
    """
    Extract content from any supported file.

    Auto-detects file type and uses appropriate extractor.

    Args:
        file_path: Path to file
        config: Extraction configuration

    Returns:
        ExtractionResult

    Example:
        result = extract_file("document.pdf")
        print(result.text)
    """
    file_path = Path(file_path)
    extractor = create_extractor(file_path, config)
    return extractor.extract(file_path, config)
