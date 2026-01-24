"""
Document Generator für Explosionsschutzdokumente
================================================

Generiert Word-Dokumente aus Session-Daten mit:
- Corporate Design (Template/Logo)
- Strukturierte Kapitel (Phasen 1-13)
- Inhaltsverzeichnis
- Anhänge (optional)
"""

import io
import os
from datetime import datetime
from typing import List, Optional, Dict, Any

from django.conf import settings

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from apps.expert_hub.models import (
    ExAnalysisSession,
    ExSessionPhaseStatus,
    ExWorkflowPhase,
    ExSessionDocument,
)


class ExSchutzDocumentGenerator:
    """Generiert Explosionsschutzdokumente als Word-Datei."""
    
    def __init__(self, session: ExAnalysisSession):
        self.session = session
        self.document = None
        self.phases = []
        self.phase_statuses = {}
        
    def load_data(self):
        """Lädt alle Phasen-Daten für die Session."""
        # Alle Top-Level Phasen
        self.phases = ExWorkflowPhase.objects.filter(
            parent__isnull=True
        ).order_by('order')
        
        # Status für jede Phase
        for phase in self.phases:
            status = ExSessionPhaseStatus.objects.filter(
                session=self.session,
                phase=phase
            ).first()
            if status:
                self.phase_statuses[phase.id] = status
                
            # Auch Unter-Phasen laden
            for child in phase.children.all():
                child_status = ExSessionPhaseStatus.objects.filter(
                    session=self.session,
                    phase=child
                ).first()
                if child_status:
                    self.phase_statuses[child.id] = child_status
    
    def create_document(self, template_path: Optional[str] = None) -> Document:
        """
        Erstellt das Word-Dokument.
        
        Args:
            template_path: Pfad zur Word-Vorlage (optional)
            
        Returns:
            Document: python-docx Document Objekt
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx ist nicht installiert. Bitte 'pip install python-docx' ausführen.")
        
        # Template oder leeres Dokument
        if template_path and os.path.exists(template_path):
            self.document = Document(template_path)
        elif self.session.template_file:
            self.document = Document(self.session.template_file.path)
        else:
            self.document = Document()
            self._setup_default_styles()
        
        self.load_data()
        
        # Dokument aufbauen
        self._add_cover_page()
        self._add_table_of_contents()
        self._add_content_sections()
        self._add_appendix_section()
        self._add_revision_history()
        
        return self.document
    
    def _setup_default_styles(self):
        """Richtet Standard-Styles ein wenn keine Vorlage verwendet wird."""
        styles = self.document.styles
        
        # Überschrift 1
        if 'Heading 1' in styles:
            h1 = styles['Heading 1']
            h1.font.size = Pt(16)
            h1.font.bold = True
            
        # Überschrift 2
        if 'Heading 2' in styles:
            h2 = styles['Heading 2']
            h2.font.size = Pt(14)
            h2.font.bold = True
    
    def _add_cover_page(self):
        """Fügt Deckblatt hinzu."""
        # Logo (falls vorhanden)
        if self.session.company_logo:
            try:
                self.document.add_picture(
                    self.session.company_logo.path,
                    width=Inches(2)
                )
            except Exception:
                pass  # Logo konnte nicht geladen werden
        
        # Titel
        self.document.add_paragraph()
        self.document.add_paragraph()
        
        title = self.document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("EXPLOSIONSSCHUTZDOKUMENT")
        run.bold = True
        run.font.size = Pt(24)
        
        # Untertitel
        subtitle = self.document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("gemäß § 6 Abs. 9 GefStoffV")
        run.font.size = Pt(14)
        
        self.document.add_paragraph()
        self.document.add_paragraph()
        
        # Projekt-Info
        info = self.document.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.add_run(f"Projekt: {self.session.project_name or self.session.name}").bold = True
        
        if self.session.project_location:
            loc = self.document.add_paragraph()
            loc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            loc.add_run(f"Standort: {self.session.project_location}")
        
        self.document.add_paragraph()
        self.document.add_paragraph()
        self.document.add_paragraph()
        
        # Meta-Info Tabelle
        table = self.document.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        meta_data = [
            ("Dokument-Nr.:", self.session.metadata.get('doc_number', f"EX-{self.session.id.hex[:8].upper()}")),
            ("Version:", self.session.metadata.get('version', '1.0')),
            ("Erstellt am:", datetime.now().strftime('%d.%m.%Y')),
            ("Erstellt von:", self.session.created_by.get_full_name() if self.session.created_by else 'N/A'),
        ]
        
        for i, (label, value) in enumerate(meta_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)
        
        # Seitenumbruch
        self.document.add_page_break()
    
    def _add_table_of_contents(self):
        """Fügt Inhaltsverzeichnis hinzu (als Platzhalter für Word)."""
        self.document.add_heading("Inhaltsverzeichnis", level=1)
        
        # TOC-Feld einfügen (wird in Word aktualisiert)
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()
        
        # Word TOC Feldcode
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)
        
        # Hinweis
        hint = self.document.add_paragraph()
        hint.add_run("(Bitte Inhaltsverzeichnis in Word aktualisieren: Rechtsklick → Felder aktualisieren)").italic = True
        
        self.document.add_page_break()
    
    def _add_content_sections(self):
        """Fügt alle Kapitel/Phasen hinzu."""
        for phase in self.phases:
            self._add_phase_section(phase, level=1)
            
            # Unter-Phasen
            for child in phase.children.all().order_by('order'):
                self._add_phase_section(child, level=2)
                
                # Sub-Sub-Phasen
                for subchild in child.children.all().order_by('order'):
                    self._add_phase_section(subchild, level=3)
    
    def _add_phase_section(self, phase: ExWorkflowPhase, level: int = 1):
        """Fügt einen Phasen-Abschnitt hinzu."""
        # Überschrift
        self.document.add_heading(f"{phase.number} {phase.title}", level=level)
        
        # Status holen
        status = self.phase_statuses.get(phase.id)
        
        if status and status.content:
            # Inhalt einfügen (Markdown zu Text)
            content = status.content.strip()
            
            # Einfache Markdown-Verarbeitung
            for line in content.split('\n'):
                line = line.strip()
                if not line:
                    continue
                    
                # Überschriften ignorieren (bereits als Word-Heading)
                if line.startswith('##'):
                    continue
                    
                # Listen
                if line.startswith('- ') or line.startswith('* '):
                    p = self.document.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('1. ') or line.startswith('2. '):
                    p = self.document.add_paragraph(line[3:], style='List Number')
                # Tabellen (vereinfacht)
                elif line.startswith('|'):
                    # Tabellen werden als Text eingefügt
                    self.document.add_paragraph(line)
                else:
                    # Normaler Text
                    # Bold-Text verarbeiten
                    if '**' in line:
                        p = self.document.add_paragraph()
                        parts = line.split('**')
                        for i, part in enumerate(parts):
                            if part:
                                run = p.add_run(part)
                                if i % 2 == 1:  # Ungerade = bold
                                    run.bold = True
                    else:
                        self.document.add_paragraph(line)
        else:
            # Platzhalter für leere Phasen
            p = self.document.add_paragraph()
            run = p.add_run("[Noch nicht bearbeitet]")
            run.italic = True
            run.font.color.rgb = None  # Grau
    
    def _add_appendix_section(self):
        """Fügt Anhang-Abschnitt hinzu."""
        self.document.add_page_break()
        self.document.add_heading("Anhänge", level=1)
        
        # Dokumente auflisten
        documents = ExSessionDocument.objects.filter(session=self.session)
        
        if documents.exists():
            self.document.add_heading("A. Verzeichnis der Anlagen", level=2)
            
            table = self.document.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            # Header
            headers = table.rows[0].cells
            headers[0].text = "Nr."
            headers[1].text = "Dokument"
            headers[2].text = "Typ"
            
            for i, doc in enumerate(documents, 1):
                row = table.add_row().cells
                row[0].text = f"A.{i}"
                row[1].text = doc.original_filename
                row[2].text = doc.get_document_type_display()
        else:
            self.document.add_paragraph("Keine Anhänge vorhanden.")
    
    def _add_revision_history(self):
        """Fügt Revisionsverlauf hinzu."""
        self.document.add_page_break()
        self.document.add_heading("Revisionsverlauf", level=1)
        
        table = self.document.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        
        # Header
        headers = table.rows[0].cells
        headers[0].text = "Version"
        headers[1].text = "Datum"
        headers[2].text = "Autor"
        headers[3].text = "Änderungen"
        
        # Erste Version
        row = table.rows[1].cells
        row[0].text = "1.0"
        row[1].text = datetime.now().strftime('%d.%m.%Y')
        row[2].text = self.session.created_by.get_full_name() if self.session.created_by else 'N/A'
        row[3].text = "Erstversion"
    
    def get_html_preview(self) -> str:
        """Generiert HTML-Vorschau des Dokuments."""
        self.load_data()
        
        html_parts = []
        
        # CSS
        html_parts.append("""
        <style>
            .ex-doc { font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; }
            .ex-doc h1 { color: #0d6efd; border-bottom: 2px solid #0d6efd; padding-bottom: 10px; }
            .ex-doc h2 { color: #333; margin-top: 30px; }
            .ex-doc h3 { color: #666; }
            .ex-doc .cover { text-align: center; padding: 50px 0; border: 1px solid #ddd; margin-bottom: 30px; }
            .ex-doc .meta-table { margin: 20px auto; border-collapse: collapse; }
            .ex-doc .meta-table td { border: 1px solid #ddd; padding: 8px 15px; }
            .ex-doc .phase { margin: 20px 0; padding: 15px; background: #f8f9fa; border-radius: 5px; }
            .ex-doc .phase.empty { background: #fff3cd; }
            .ex-doc .content { white-space: pre-wrap; }
            .ex-doc .toc { background: #e9ecef; padding: 20px; margin: 20px 0; }
            .ex-doc .toc ul { list-style: none; padding-left: 20px; }
            .ex-doc .appendix { margin-top: 40px; padding-top: 20px; border-top: 2px solid #ddd; }
        </style>
        """)
        
        html_parts.append('<div class="ex-doc">')
        
        # Deckblatt
        html_parts.append('<div class="cover">')
        if self.session.company_logo:
            html_parts.append(f'<img src="{self.session.company_logo.url}" style="max-width: 150px;"><br><br>')
        html_parts.append('<h1>EXPLOSIONSSCHUTZDOKUMENT</h1>')
        html_parts.append('<p>gemäß § 6 Abs. 9 GefStoffV</p>')
        html_parts.append(f'<h2>{self.session.project_name or self.session.name}</h2>')
        if self.session.project_location:
            html_parts.append(f'<p>{self.session.project_location}</p>')
        html_parts.append('<table class="meta-table">')
        html_parts.append(f'<tr><td>Dokument-Nr.:</td><td>EX-{self.session.id.hex[:8].upper()}</td></tr>')
        html_parts.append(f'<tr><td>Version:</td><td>1.0</td></tr>')
        html_parts.append(f'<tr><td>Erstellt am:</td><td>{datetime.now().strftime("%d.%m.%Y")}</td></tr>')
        html_parts.append('</table>')
        html_parts.append('</div>')
        
        # Inhaltsverzeichnis
        html_parts.append('<div class="toc">')
        html_parts.append('<h2>Inhaltsverzeichnis</h2>')
        html_parts.append('<ul>')
        for phase in self.phases:
            status = self.phase_statuses.get(phase.id)
            icon = '✓' if status and status.content else '○'
            html_parts.append(f'<li>{icon} <a href="#phase-{phase.id}">{phase.number} {phase.title}</a></li>')
            
            if phase.children.exists():
                html_parts.append('<ul>')
                for child in phase.children.all().order_by('order'):
                    child_status = self.phase_statuses.get(child.id)
                    child_icon = '✓' if child_status and child_status.content else '○'
                    html_parts.append(f'<li>{child_icon} {child.number} {child.title}</li>')
                html_parts.append('</ul>')
        html_parts.append('</ul>')
        html_parts.append('</div>')
        
        # Kapitel
        for phase in self.phases:
            status = self.phase_statuses.get(phase.id)
            has_content = status and status.content
            
            html_parts.append(f'<div class="phase {"" if has_content else "empty"}" id="phase-{phase.id}">')
            html_parts.append(f'<h2>{phase.number} {phase.title}</h2>')
            
            if has_content:
                # Markdown zu HTML (vereinfacht)
                content = status.content
                content = content.replace('**', '<strong>').replace('**', '</strong>')
                content = content.replace('\n', '<br>')
                html_parts.append(f'<div class="content">{content}</div>')
            else:
                html_parts.append('<p><em>[Noch nicht bearbeitet]</em></p>')
            
            html_parts.append('</div>')
            
            # Unter-Phasen
            for child in phase.children.all().order_by('order'):
                child_status = self.phase_statuses.get(child.id)
                has_child_content = child_status and child_status.content
                
                html_parts.append(f'<div class="phase {"" if has_child_content else "empty"}">')
                html_parts.append(f'<h3>{child.number} {child.title}</h3>')
                
                if has_child_content:
                    content = child_status.content
                    content = content.replace('\n', '<br>')
                    html_parts.append(f'<div class="content">{content}</div>')
                else:
                    html_parts.append('<p><em>[Noch nicht bearbeitet]</em></p>')
                
                html_parts.append('</div>')
        
        # Anhänge
        documents = ExSessionDocument.objects.filter(session=self.session)
        if documents.exists():
            html_parts.append('<div class="appendix">')
            html_parts.append('<h2>Anhänge</h2>')
            html_parts.append('<ul>')
            for i, doc in enumerate(documents, 1):
                html_parts.append(f'<li>A.{i}: {doc.original_filename} ({doc.get_document_type_display()})</li>')
            html_parts.append('</ul>')
            html_parts.append('</div>')
        
        html_parts.append('</div>')
        
        return '\n'.join(html_parts)
    
    def save_to_buffer(self) -> io.BytesIO:
        """Speichert Dokument in BytesIO Buffer für Download."""
        if not self.document:
            self.create_document()
            
        buffer = io.BytesIO()
        self.document.save(buffer)
        buffer.seek(0)
        return buffer
    
    def get_filename(self) -> str:
        """Generiert Dateinamen für Download."""
        project = self.session.project_name or self.session.name
        project = "".join(c for c in project if c.isalnum() or c in (' ', '-', '_')).strip()
        project = project.replace(' ', '_')
        date = datetime.now().strftime('%Y%m%d')
        return f"Explosionsschutzdokument_{project}_{date}.docx"


def generate_exschutz_document(session_id: str, include_attachments: List[str] = None) -> io.BytesIO:
    """
    Convenience-Funktion zum Generieren eines Explosionsschutzdokuments.
    
    Args:
        session_id: UUID der Session
        include_attachments: Liste von Dokument-IDs für Anhänge
        
    Returns:
        BytesIO Buffer mit Word-Dokument
    """
    session = ExAnalysisSession.objects.get(id=session_id)
    generator = ExSchutzDocumentGenerator(session)
    generator.create_document()
    return generator.save_to_buffer()
