"""
Research Export Service
=======================

Export research results to various formats:
- DOCX (Word)
- LaTeX
- Markdown
- PDF (via LaTeX or HTML)
"""

import logging
import io
from typing import Dict, List, Optional, Any
from datetime import datetime

logger = logging.getLogger(__name__)


class ResearchExportService:
    """
    Export research projects to various document formats.
    
    Usage:
        service = ResearchExportService()
        docx_bytes = service.export_docx(project)
        latex_str = service.export_latex(project)
    """
    
    def export_markdown(self, project, include_sources: bool = True) -> str:
        """
        Export project to Markdown format.
        
        Args:
            project: ResearchProject instance
            include_sources: Include bibliography
            
        Returns:
            Markdown string
        """
        lines = []
        
        # Title
        lines.append(f"# {project.name}")
        lines.append("")
        lines.append(f"*Research Type: {project.get_research_type_display()}*")
        lines.append(f"*Created: {project.created_at.strftime('%Y-%m-%d')}*")
        lines.append("")
        
        # Description
        if project.description:
            lines.append("## Overview")
            lines.append("")
            lines.append(project.description)
            lines.append("")
        
        # Query
        lines.append("## Research Query")
        lines.append("")
        lines.append(f"> {project.query}")
        lines.append("")
        
        # Findings
        findings = project.findings.all()
        if findings.exists():
            lines.append("## Key Findings")
            lines.append("")
            for i, finding in enumerate(findings, 1):
                lines.append(f"### Finding {i}: {finding.title}")
                lines.append("")
                lines.append(finding.content)
                if finding.confidence_score:
                    lines.append(f"\n*Confidence: {finding.confidence_score:.0%}*")
                lines.append("")
        
        # Sources
        if include_sources:
            sources = project.sources.all()
            if sources.exists():
                lines.append("## Sources")
                lines.append("")
                for i, source in enumerate(sources, 1):
                    lines.append(f"{i}. [{source.title}]({source.url})")
                    if source.snippet:
                        lines.append(f"   > {source.snippet[:200]}...")
                    lines.append("")
        
        return "\n".join(lines)
    
    def export_latex(
        self,
        project,
        include_sources: bool = True,
        citation_style: str = 'apa'
    ) -> str:
        """
        Export project to LaTeX format.
        
        Args:
            project: ResearchProject instance
            include_sources: Include bibliography
            citation_style: Citation style (apa, mla, ieee)
            
        Returns:
            LaTeX document string
        """
        # Get citation style from project if available
        style = getattr(project, 'citation_style', citation_style) or citation_style
        
        lines = [
            r"\documentclass[12pt,a4paper]{article}",
            r"\usepackage[utf8]{inputenc}",
            r"\usepackage[T1]{fontenc}",
            r"\usepackage{hyperref}",
            r"\usepackage{booktabs}",
            r"\usepackage{graphicx}",
            f"\\usepackage[style={style}]{{biblatex}}",
            r"\addbibresource{references.bib}",
            "",
            f"\\title{{{self._escape_latex(project.name)}}}",
            r"\author{BF Agent Research Hub}",
            f"\\date{{{project.created_at.strftime('%B %d, %Y')}}}",
            "",
            r"\begin{document}",
            r"\maketitle",
            "",
        ]
        
        # Abstract / Description
        if project.description:
            lines.extend([
                r"\begin{abstract}",
                self._escape_latex(project.description),
                r"\end{abstract}",
                "",
            ])
        
        # Introduction with Query
        lines.extend([
            r"\section{Research Query}",
            self._escape_latex(project.query),
            "",
        ])
        
        # Findings
        findings = project.findings.all()
        if findings.exists():
            lines.append(r"\section{Key Findings}")
            lines.append("")
            for finding in findings:
                lines.append(f"\\subsection{{{self._escape_latex(finding.title)}}}")
                lines.append(self._escape_latex(finding.content))
                if finding.confidence_score:
                    lines.append(f"\n\\textit{{Confidence: {finding.confidence_score:.0%}}}")
                lines.append("")
        
        # Bibliography
        if include_sources:
            lines.extend([
                r"\section{References}",
                r"\printbibliography[heading=none]",
                "",
            ])
        
        lines.append(r"\end{document}")
        
        return "\n".join(lines)
    
    def export_bibtex(self, project) -> str:
        """
        Export sources as BibTeX bibliography.
        
        Args:
            project: ResearchProject instance
            
        Returns:
            BibTeX string
        """
        entries = []
        sources = project.sources.all()
        
        for i, source in enumerate(sources, 1):
            entry_id = f"source{i}"
            
            # Determine entry type
            entry_type = "online"  # Default for web sources
            
            entry = [
                f"@{entry_type}{{{entry_id},",
                f"  title = {{{self._escape_bibtex(source.title)}}},",
            ]
            
            # Add author if available
            author = source.metadata.get('author') if source.metadata else None
            if author:
                entry.append(f"  author = {{{self._escape_bibtex(author)}}},")
            
            # Add URL
            entry.append(f"  url = {{{source.url}}},")
            
            # Add access date
            access_date = source.created_at.strftime('%Y-%m-%d')
            entry.append(f"  urldate = {{{access_date}}},")
            
            # Add year if available
            year = source.metadata.get('year') if source.metadata else None
            if year:
                entry.append(f"  year = {{{year}}},")
            else:
                entry.append(f"  year = {{{source.created_at.year}}},")
            
            entry.append("}")
            entries.append("\n".join(entry))
        
        return "\n\n".join(entries)
    
    def export_docx(self, project, include_sources: bool = True) -> bytes:
        """
        Export project to DOCX (Word) format.
        
        Args:
            project: ResearchProject instance
            include_sources: Include bibliography
            
        Returns:
            DOCX file as bytes
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.warning("python-docx not installed, returning markdown as fallback")
            return self.export_markdown(project, include_sources).encode('utf-8')
        
        doc = Document()
        
        # Title
        title = doc.add_heading(project.name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        meta = doc.add_paragraph()
        meta.add_run(f"Research Type: {project.get_research_type_display()}").italic = True
        meta.add_run(f"\nCreated: {project.created_at.strftime('%Y-%m-%d')}").italic = True
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Description
        if project.description:
            doc.add_heading("Overview", level=1)
            doc.add_paragraph(project.description)
        
        # Query
        doc.add_heading("Research Query", level=1)
        query_para = doc.add_paragraph()
        query_para.add_run(project.query).italic = True
        
        # Findings
        findings = project.findings.all()
        if findings.exists():
            doc.add_heading("Key Findings", level=1)
            for finding in findings:
                doc.add_heading(finding.title, level=2)
                doc.add_paragraph(finding.content)
                if finding.confidence_score:
                    conf = doc.add_paragraph()
                    conf.add_run(f"Confidence: {finding.confidence_score:.0%}").italic = True
        
        # Sources
        if include_sources:
            sources = project.sources.all()
            if sources.exists():
                doc.add_heading("Sources", level=1)
                for i, source in enumerate(sources, 1):
                    p = doc.add_paragraph(style='List Number')
                    p.add_run(source.title).bold = True
                    p.add_run(f"\n{source.url}")
                    if source.snippet:
                        p.add_run(f"\n{source.snippet[:200]}...").italic = True
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
    def _escape_latex(self, text: str) -> str:
        """Escape special LaTeX characters."""
        if not text:
            return ""
        
        replacements = [
            ('\\', r'\textbackslash{}'),
            ('&', r'\&'),
            ('%', r'\%'),
            ('$', r'\$'),
            ('#', r'\#'),
            ('_', r'\_'),
            ('{', r'\{'),
            ('}', r'\}'),
            ('~', r'\textasciitilde{}'),
            ('^', r'\textasciicircum{}'),
        ]
        
        result = text
        for old, new in replacements:
            result = result.replace(old, new)
        
        return result
    
    def _escape_bibtex(self, text: str) -> str:
        """Escape special BibTeX characters."""
        if not text:
            return ""
        
        # BibTeX uses {} for grouping, escape them
        return text.replace('{', r'\{').replace('}', r'\}')


# Singleton instance
_export_service = None

def get_export_service() -> ResearchExportService:
    """Get singleton instance of ResearchExportService."""
    global _export_service
    if _export_service is None:
        _export_service = ResearchExportService()
    return _export_service
